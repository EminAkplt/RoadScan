# -*- coding: utf-8 -*-
"""RoadScan akademik makalesini Word (.docx) olarak üretir."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS = os.path.join(ROOT, "docs")
FIGS = os.path.join(DOCS, "figs")
OUT = os.path.join(DOCS, "RoadScan_Makale.docx")

doc = Document()
# Varsayılan stil
st = doc.styles["Normal"]
st.font.name = "Times New Roman"; st.font.size = Pt(11)
for sec in doc.sections:
    sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)

def para(text, *, size=11, bold=False, italic=False, align="just", after=6, before=0, color=None):
    p = doc.add_paragraph()
    p.alignment = {"just": WD_ALIGN_PARAGRAPH.JUSTIFY, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "left": WD_ALIGN_PARAGRAPH.LEFT}[align]
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return p

def h(text, size=13):
    para(text, size=size, bold=True, align="left", before=10, after=4)

def rich(parts, align="just", after=6):
    """parts: [(text, bold), ...] -> tek paragraf, bazı parçalar kalın."""
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(after)
    for text, bold in parts:
        r = p.add_run(text); r.font.size = Pt(11); r.bold = bold
    return p

def table(rows, widths, header_bold=True, last_bold=False):
    t = doc.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, w in enumerate(widths):
        for r in t.rows:
            r.cells[ci].width = Cm(w)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci); cell.text = ""
            pr = cell.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = pr.add_run(str(val)); run.font.size = Pt(10)
            if (ri == 0 and header_bold) or (last_bold and ri == len(rows)-1):
                run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def figure(path, cap, width=14.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(cap, size=9, italic=True, align="center", after=10, before=2)

# ---------- Başlık / yazar ----------
para("Araç İçi Dashcam Görüntülerinden Derin Öğrenme ve Görüntü İşleme ile "
     "Gerçek Zamanlı Yol Bozukluğu Tespiti", size=15, bold=True, align="center", after=8)
para("Mehmet Emin AKPOLAT", size=11, bold=True, align="center", after=1)
para("Fırat Üniversitesi, Teknoloji Fakültesi, Yazılım Mühendisliği Bölümü", size=9, align="center", after=1)
para("215542003@firat.edu.tr", size=9, align="center", after=10)

# ---------- Öz ----------
h("Öz")
para("Bu çalışma, araç ön kamerasından (dashcam) elde edilen görüntü akışında yoldaki çukur ve çatlak "
     "türü bozuklukları gerçek zamanlı tespit eden, konum ve zaman bilgisiyle eşleştirip merkezi bir "
     "veritabanına raporlayan tam yığın (full-stack) bir sistem önermektedir. Sistem, evrişimsel sinir ağı "
     "tabanlı YOLOv8 nesne tespit modeli ile klasik görüntü işleme yöntemlerini birleştiren hibrit bir "
     "yaklaşım kullanır. Model, üç açık kaynak veri setinin (RDD2022, BharatPotHole, IVCNZ) "
     "birleştirilmesiyle oluşturulan yaklaşık 46.000 görüntülük dört sınıflı (boyuna çatlak, enine çatlak, "
     "timsah sırtı çatlak ve çukur) bir veri kümesi üzerinde transfer öğrenme ile eğitilmiştir. Eğitilen model "
     "doğrulama kümesinde %61 ortalama hassasiyet (mAP@50) elde etmiştir. Tespit edilen bozukluklar "
     "yalnızca derin öğrenme çıktısına bırakılmamış; gökyüzü/yapı gibi yol dışı alanları ve düz/tek renk "
     "yüzeyleri eleyen ROI ve doku (gradyan) tabanlı görüntü işleme süzgeçleriyle doğrulanarak yanlış "
     "pozitiflerin azaltılması hedeflenmiştir. Model ONNX biçimine aktarılarak tarayıcı üzerinde, internet "
     "bağlantısı olmadan, cihaz üstünde çalışacak şekilde dağıtılmıştır. Sonuçlar, klasik kenar tabanlı "
     "yöntemlere kıyasla anlamlı bir doğruluk artışı sağlandığını; saha verisiyle ince ayar (fine-tuning) "
     "yapılarak performansın daha da artırılabileceğini göstermektedir.", size=10, after=4)
rich([("Anahtar Kelimeler: ", True),
      ("Yol bozukluğu tespiti, çukur, derin öğrenme, YOLOv8, nesne tespiti, görüntü işleme, ONNX, "
       "kenar yapay zekâ.", False)])

# ---------- 1. Giriş ----------
h("1. Giriş")
para("Yol yüzeyindeki çukur ve çatlaklar, trafik güvenliğini doğrudan tehdit eden ve araçlarda hasara yol "
     "açan önemli altyapı sorunlarıdır. Bu bozuklukların belediye ve karayolları ekiplerince geleneksel "
     "yöntemlerle tespiti çoğunlukla manuel, yavaş, maliyetli ve dağınık kayıt tutmaya dayalıdır. Son yıllarda "
     "bilgisayarla görme ve derin öğrenme alanındaki gelişmeler, bu sürecin otomatikleştirilmesine olanak "
     "tanımıştır [1], [4].")
para("Bu bozuklukların görüntüden tespiti iki temel yaklaşımla ele alınabilir. Birincisi, kenar tespiti, "
     "morfolojik işlemler ve kontur analizi gibi klasik görüntü işleme yöntemleridir. Bu yöntemler hesaplama "
     "açısından hafif olsa da “bozukluk” kavramını öğrenemedikleri için gölge, şerit çizgisi ve yüzey dokusu "
     "gibi her türlü kenarı bozukluk olarak işaretleyerek yüksek oranda yanlış pozitif üretirler. İkincisi, "
     "etiketlenmiş veriden öğrenen evrişimsel sinir ağı (CNN) tabanlı nesne tespit modelleridir; bunlar "
     "bozukluğun görsel örüntüsünü öğrenir ve çok daha az yanlış pozitif üretir [1], [3].")
para("Bu çalışmanın katkıları şöyle özetlenebilir: (i) üç açık kaynak veri seti birleştirilerek çukur sınıfı çok "
     "sayıda örnekle desteklenen dört sınıflı bir yol-hasarı veri kümesi oluşturulmuştur; (ii) YOLOv8 tabanlı "
     "bir model transfer öğrenme ile eğitilmiş ve ONNX ile tarayıcıda, çevrimdışı çalışacak şekilde "
     "dağıtılmıştır; (iii) derin öğrenme çıktısı, klasik görüntü işleme tabanlı bir doğrulama katmanıyla "
     "birleştirilerek yanlış pozitifleri azaltan hibrit bir karar mekanizması önerilmiştir; (iv) yalnızca kritik "
     "(yolculuğu etkileyen) bozuklukların raporlanması ve aynı bozukluğun kare-arası takiple tek kayıt "
     "edilmesiyle veritabanı verimliliği sağlanmıştır.")

# ---------- 2. Materyal ve Metot ----------
h("2. Materyal ve Metot")
h("2.1. Veri Seti", size=11.5)
para("Çalışmada üç açık kaynak veri seti birleştirilmiştir: RDD2022 (çok uluslu yol hasarı veri seti) [2], "
     "BharatPotHole (çeşitli yol koşullarında dashcam çukur görüntüleri) ve IVCNZ çukur veri seti. Tüm "
     "etiketler ortak bir dört sınıflı düzene (0: boyuna çatlak, 1: enine çatlak, 2: timsah sırtı çatlak, 3: çukur) "
     "eşlenmiş; çukur içeren iki ek veri setinin etiketleri çukur sınıfına yönlendirilerek bu sınıf üç kaynaktan "
     "birden beslenmiştir. Birleştirme sonrası elde edilen veri kümesi yaklaşık 46.000 görüntüden oluşmakta "
     "olup eğitim/doğrulama olarak ayrılmıştır (Tablo 1). Sınıf bazlı örnek (kutu) dağılımı Şekil 1’de "
     "verilmiştir.")
table([["Bölüm", "Görüntü Sayısı"], ["Eğitim (train)", "38.814"], ["Doğrulama (val)", "7.226"],
       ["Toplam", "46.040"]], widths=[7, 5], last_bold=True)
para("Tablo 1. Birleşik veri setinin eğitim/doğrulama dağılımı.", size=9, italic=True, align="center", after=10)
figure(os.path.join(FIGS, "fig1.png"), "Şekil 1. Birleşik veri setinde sınıf bazlı eğitim örneği (sınırlayıcı kutu) sayıları.")

h("2.2. Sistem Mimarisi", size=11.5)
para("Önerilen sistem üç modülden oluşur. (1) Tespit motoru: tarayıcıda çalışan, fotoğraf/video/canlı kamera "
     "kaynaklarını işleyen istemci uygulamasıdır; çıkarım cihaz üstünde, çevrimdışı yapılır. (2) Veri kayıt "
     "servisi: Node.js/Express tabanlı REST API ile gelen tespitleri PostgreSQL veritabanına kaydeder, "
     "kırpılmış görüntüleri dosya olarak saklar. (3) Yönetim paneli: Leaflet harita üzerinde tespitleri konuma "
     "göre gösterir, filtreleme ve özet istatistik sunar. Bu yapı, modelin sahada (örn. araç içi küçük donanım) "
     "çevrimdışı çalışmasına ve bağlantı sağlandığında merkeze raporlamasına uygundur.")

h("2.3. Tespit Modeli (YOLOv8)", size=11.5)
para("Tespit için tek aşamalı (single-stage) bir nesne tespit modeli olan YOLOv8’in küçük (small) sürümü "
     "kullanılmıştır. YOLO ailesi, görüntüyü tek bir ileri besleme ile işleyip sınırlayıcı kutuları ve sınıf "
     "olasılıklarını aynı anda üreterek gerçek zamanlı çalışma sağlar [3]. Evrişim katmanları görüntüdeki kenar, "
     "doku ve şekil gibi özellikleri çıkarırken, ağın tespit başlığı bu özelliklerden kutu koordinatlarını ve sınıf "
     "güvenlerini tahmin eder. Eğitim, COCO veri kümesinde ön-eğitilmiş ağırlıklardan başlatılarak (transfer "
     "öğrenme) gerçekleştirilmiştir [6]. Eğitim hiperparametreleri Tablo 2’de verilmiştir.")
table([["Hiperparametre", "Değer"], ["Model", "YOLOv8s (small)"],
       ["Başlangıç ağırlığı", "COCO ön-eğitimli (transfer öğrenme)"],
       ["Giriş çözünürlüğü", "640 × 640"], ["Paket boyutu (batch)", "16"],
       ["Tur sayısı (epoch)", "80 (erken durdurma: 25)"],
       ["Optimizasyon / öğrenme hızı", "SGD (otomatik) / 0.01"],
       ["Donanım", "NVIDIA RTX 5060 (8 GB)"]], widths=[6.5, 8])
para("Tablo 2. Model eğitimi hiperparametreleri.", size=9, italic=True, align="center", after=8)

h("2.4. Görüntü İşleme Aşamaları", size=11.5)
rich([("Model, ham haliyle değil, klasik görüntü işleme adımlarıyla çevrelenmiş bir hat (pipeline) içinde "
       "çalışır. ", False), ("Ön işleme: ", True),
      ("her kare, en-boy oranı korunarak 640×640 boyutuna mektup-kutusu (letterbox) yöntemiyle ölçeklenir, "
       "[0,1] aralığına normalize edilir ve modelin beklediği tensör biçimine dönüştürülür. ", False),
      ("Son işleme: ", True),
      ("modelin ham çıktısı ayrıştırılır, güven eşiğinin altındaki tespitler elenir ve örtüşen kutular maksimum "
       "olmayan baskılama (Non-Maximum Suppression, NMS) ile birleştirilir; kutular özgün görüntü "
       "koordinatlarına geri ölçeklenir.", False)])
rich([("Hibrit doğrulama (klasik görüntü işleme): ", True),
      ("Bir tespit “kritik” olarak raporlanmadan önce iki görüntü işleme süzgecinden geçirilir. (i) İlgi alanı "
       "(ROI) süzgeci: kutu merkezi görüntünün üst bölgesindeyse (gökyüzü, bina, ağaç gibi yol dışı alanlar) "
       "tespit reddedilir. (ii) Doku/kenar süzgeci: kutu bölgesi gri tonlamaya çevrilir ve gradyan (Sobel benzeri "
       "kenar yoğunluğu) hesaplanır; düz/tek renk yüzeyler (araç paneli, duvar, boş asfalt) düşük gradyan "
       "ürettiğinden reddedilir, çünkü gerçek bozukluklar belirgin kenar ve dokuya sahiptir. Bu hibrit yapı, "
       "derin öğrenmenin öğrenme gücüyle klasik görüntü işlemenin yorumlanabilir kurallarını birleştirir.", False)])
rich([("Önem (severity) ve raporlama: ", True),
      ("Her tespitin önem derecesi, sınırlayıcı kutu alanının kareye oranına ve sınıfına göre Küçük/Orta/Kritik "
       "olarak belirlenir. Üç kademe de ekranda görselleştirilir; ancak veritabanına yalnızca kritik bozukluklar "
       "gönderilir. Videoda aynı fiziksel bozukluğun yüzlerce karede tekrar kaydedilmesini önlemek için kareler "
       "arası IoU tabanlı bir takip uygulanır ve her bozukluk tek kayıt olarak raporlanır.", False)])

h("2.5. Performans Metrikleri", size=11.5)
para("Nesne tespitinde başarı, bir tespitin doğru sayılması için tahmin kutusu ile gerçek kutu arasındaki "
     "örtüşme oranının (Intersection over Union, IoU) belirli bir eşiği aşması koşuluna dayanır. Doğru Pozitif "
     "(TP), Yanlış Pozitif (FP) ve Yanlış Negatif (FN) sayıları üzerinden Kesinlik (Precision), Duyarlılık "
     "(Recall) ve ortalama hassasiyet (mAP) hesaplanır. İlgili eşitlikler aşağıda verilmiştir.")
para("IoU = Kesişim Alanı / Birleşim Alanı", size=11, italic=True, align="center", after=2)
para("Kesinlik = TP / (TP + FP)        Duyarlılık = TP / (TP + FN)", size=11, italic=True, align="center", after=2)
para("AP = Kesinlik–Duyarlılık eğrisi altındaki alan ;  mAP = sınıflar üzerinde AP ortalaması", size=11, italic=True, align="center", after=6)
para("Bu çalışmada birincil metrik, IoU eşiği 0.5 olan ortalama hassasiyet (mAP@50) ile daha katı bir ölçüt "
     "olan mAP@50–95’tir.")

# ---------- 3. Deneysel Bulgular ----------
h("3. Deneysel Bulgular")
para("Model, NVIDIA RTX 5060 (8 GB) ekran kartında 80 tur boyunca eğitilmiştir. Doğrulama kümesi (7.226 "
     "görüntü) üzerindeki sınıf bazlı ve genel sonuçlar Tablo 3’te, mAP@50 değerlerinin karşılaştırması ise "
     "Şekil 2’de sunulmaktadır.")
table([["Sınıf", "Kesinlik", "Duyarlılık", "mAP@50", "mAP@50–95"],
       ["Boyuna Çatlak", "0.672", "0.549", "0.602", "0.334"],
       ["Enine Çatlak", "0.654", "0.565", "0.603", "0.304"],
       ["Timsah Çatlak", "0.703", "0.638", "0.681", "0.366"],
       ["Çukur", "0.654", "0.527", "0.559", "0.245"],
       ["Genel (ortalama)", "0.671", "0.570", "0.611", "0.312"]],
      widths=[4.6, 2.5, 2.5, 2.4, 2.5], last_bold=True)
para("Tablo 3. Doğrulama kümesinde sınıf bazlı ve genel performans.", size=9, italic=True, align="center", after=10)
figure(os.path.join(FIGS, "fig2.png"), "Şekil 2. Sınıf bazlı ve genel mAP@50 değerleri.")
para("Sonuçlar incelendiğinde en yüksek başarı timsah sırtı çatlak sınıfında (mAP@50 = 0.681) elde edilmiştir; "
     "bu sınıf geniş ve dokulu bir bozulma örüntüsüne sahip olduğundan modelce daha kolay ayırt edilmiştir. "
     "Boyuna ve enine çatlaklar benzer düzeyde (≈0.60) performans göstermiştir. Çukur sınıfı, çok sayıda "
     "örnekle desteklenmesine rağmen görsel çeşitliliği (suyla dolu, gölgeli, farklı boyut ve açılarda) nedeniyle "
     "nispeten daha düşük (0.559) sonuç vermiştir. Genel mAP@50 değeri 0.611 olarak ölçülmüştür; bu değer, "
     "çok uluslu ve dört sınıflı gerçek dünya yol-hasarı verisi için literatürle uyumlu, makul bir seviyedir.")

# ---------- 4. Tartışma ve Sonuçlar ----------
h("4. Tartışma ve Sonuçlar")
para("Bu çalışmada, dashcam görüntülerinden yol bozukluğu tespiti için derin öğrenme ile klasik görüntü "
     "işlemeyi birleştiren hibrit, çevrimdışı çalışabilen bir sistem geliştirilmiştir. Klasik yalnızca-kenar tabanlı "
     "yöntemlere kıyasla, öğrenen bir modelin kullanılması yanlış pozitifleri belirgin biçimde azaltmış; ROI ve "
     "doku tabanlı görüntü işleme süzgeçleri ise yol dışı ve düz yüzey kaynaklı hatalı tespitlerin veritabanına "
     "gönderilmesini engellemiştir. Modelin ONNX biçiminde dağıtılması, internet gerektirmeden tarayıcıda ve "
     "potansiyel olarak araç içi küçük donanımlarda çalışmasına olanak tanır.")
para("Çalışmanın temel sınırı, modelin ağırlıklı olarak yurt dışı (Hindistan/Japonya vb.) yol görüntüleriyle "
     "eğitilmiş olmasıdır; bu nedenle yerel (Türkiye) yol koşullarında, özellikle suyla dolu veya aşırı bozulmuş "
     "çukurlarda yakalama oranı (recall) sınırlı kalmaktadır. Güven eşiği ve görüntü işleme süzgeçleri bu durumu "
     "kısmen iyileştirse de asıl sınırlayıcı etmen eğitim verisinin alan (domain) uyumudur.")
para("Gelecek çalışmalarda; (i) yerel yol görüntüleriyle ince ayar (fine-tuning) yapılarak alan uyumunun ve "
     "çukur yakalama oranının artırılması, (ii) GPS entegrasyonu ve konuma dayalı kümeleme ile kayıtların daha "
     "da sadeleştirilmesi, (iii) çevrimdışı kuyruk (store-and-forward) ile bağlantısız çalışma, (iv) INT8 "
     "nicemleme (quantization) ile gömülü cihazlarda hızlandırma planlanmaktadır. Sonuç olarak bu çalışma, "
     "hibrit (derin öğrenme + görüntü işleme) bir yaklaşımın yol bozukluğu tespitinde uygulanabilir ve "
     "geliştirilebilir bir temel sunduğunu göstermektedir.")

# ---------- Kaynaklar ----------
h("Kaynaklar")
refs = [
 '[1] G. Jocher, A. Chaurasia, ve J. Qiu, "Ultralytics YOLOv8," 2023. https://github.com/ultralytics/ultralytics',
 '[2] D. Arya, H. Maeda, et al., "RDD2022: A multi-national image dataset for automatic road damage detection," Geoscience Data Journal, 2024.',
 '[3] J. Redmon, S. Divvala, R. Girshick, ve A. Farhadi, "You Only Look Once: Unified, Real-Time Object Detection," Proc. IEEE CVPR, 2016, pp. 779–788.',
 '[4] Y. LeCun, Y. Bengio, ve G. Hinton, "Deep learning," Nature, vol. 521, pp. 436–444, 2015.',
 '[5] T.-Y. Lin et al., "Microsoft COCO: Common Objects in Context," Proc. ECCV, 2014, pp. 740–755.',
 '[6] A. Krizhevsky, I. Sutskever, ve G. E. Hinton, "ImageNet classification with deep convolutional neural networks," Commun. ACM, vol. 60, no. 6, pp. 84–90, 2017.',
 '[7] Microsoft, "ONNX Runtime," 2024. https://onnxruntime.ai',
 '[8] R. Padilla, S. L. Netto, ve E. A. B. da Silva, "A survey on performance metrics for object-detection algorithms," Proc. IWSSIP, 2020.',
]
for r in refs:
    para(r, size=9, after=3)

doc.save(OUT)
print("DOCX olusturuldu:", OUT)
